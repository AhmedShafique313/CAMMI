import os
from openai import OpenAI
from dotenv import load_dotenv
from docx import Document
import re

load_dotenv(dotenv_path=r"C:\Users\Kavtech AI Engineer\Documents\CAMMI\.env")
client = OpenAI(
    base_url="https://router.huggingface.co/v1",
    api_key=os.environ["HF_TOKEN"],
)

company_name = input("Enter company name: ")
industry_market_name = input("Enter industry/market name: ")
core_value_promise = input("Enter core value promise: ")
solution_platform_name = input("Enter solution/platform name: ")
target_industries_geos_companySizes = input("Enter target industries, geos, company sizes: ")
awards_recognitions = input("Enter awards and recognitions: ")
customer_logos = input("Enter customer names/logos: ")
business_concept = input("Enter the business concept: ")

system_prompt = f"""
Prompt1: You are a senior B2B marketing strategist. Given the industry: [Industry / Market Name] {industry_market_name}, write two sections in the same style and tone as a high-end B2B messaging document: Industry Challenge – Describe the major obstacles organizations face in achieving predictable revenue growth in this industry. Keep it direct and compelling. Industry Shift – Explain how buyer behavior, technology, and competitive dynamics are changing in this industry, and why traditional tactics are failing. Include 3–5 relevant, credible-sounding industry statistics, each with a short explanation in parentheses, to match professional marketing copy style. After listing the shifts and statistics, add a bold subheading: As a result of these shifts: and follow with 3–4 short bullet points summarizing the key consequences for businesses in this industry. These bullets should directly connect the shifts to the challenges they cause.  

Prompt2: Using the core value promise: {core_value_promise}, and the challenges from the Industry Challenge/Shift sections, write: Our Response – Position the company as the clear solution to the industry’s biggest problems, making a direct link to the specific issues identified in the challenges. Our Promise – Summarize in one concise paragraph how the company delivers measurable results, directly tying back to the value promise and the pain points. Opening 3–4 paragraphs of Messaging Overview – Present a compelling narrative that: Starts with a strong transition from Our Promise. Repeats the key industry shifts and challenges in context. Explains the stakes of not adapting. Teases the company’s unique approach without listing full capabilities yet. Keep the tone persuasive, strategic, and aligned with the style of a high-end B2B messaging document.  

Prompt3: The company offers a solution called [Solution / Platform Name] {solution_platform_name}. Write: Our Solution – 1–2 paragraphs introducing the platform, its role, and its impact on the target market. A bold subheading titled The [Solution / Platform Name] Platform – One short line summarizing what it is. Capabilities We Provide – A bullet list of 8–12 realistic capabilities relevant to the industry, written as bolded feature names followed by a short benefit-focused description (1 sentence each). Keep tone confident and benefits-oriented, matching the style of a high-end B2B messaging document.  

Prompt4: The target market is: [Target Industries / Geos / Company Sizes] {target_industries_geos_companySizes}. Write the How We Do It section with 4 bolded pillars: UNCOVER DEMAND, PRIORITIZE ACTIONS, ENGAGE BUYING TEAMS, and MEASURE RESULTS. For each pillar, include 4 sub-points, each starting with a bold short phrase followed by a sentence explaining the tactic. Each tactic must connect to the target market’s sales/marketing realities and challenges from the Industry Challenge section. Make sure the bolded pillar headings are written in ALL CAPS and appear exactly in this order. Format the sub-points exactly like a professional marketing playbook.  

Prompt5: Using these awards and recognitions: {awards_recognitions}, write: The [Company Name] {company_name} Difference — What Sets Us Apart – Present six bolded subheadings, each representing a core differentiator. For each subheading, include 2–3 short, benefit-focused points or sentences explaining the differentiator in detail, in the tone and style of a high-end B2B marketing document. The six subheadings must be: Rich [Industry/Platform] Data Patented Account Identification (or equivalent IP-related capability) Patented AI Predictions Deep [Customer/Operational] Insights Cross-Channel Activation One Platform for [Core Teams/Functions] Where relevant, weave in the provided awards and recognitions to strengthen credibility. End the section with one strong sentence summarizing why the company is the leader in its category. 

Prompt6: Using these customer names/logos: {customer_logos}, write: Customer Proven – Showcase social proof with these clients. If no real metrics are provided, create believable, industry-relevant proof points (e.g., “40% more opportunities, 2X win rate”). Referenceable Customers – Provide a separate bolded heading and list or describe 3–6 notable customers that can be referenced in marketing, in a style consistent with a high-end B2B messaging document. Include customer logos/names where provided. About Us Boilerplate – Provide 3 versions: 100-word 50-word One-line Each should clearly describe the company, market focus, solution, and differentiators. Press Release Version – PR-ready description including a short leadership quote (CEO/Founder style) that reinforces the value proposition and credibility. The Narrative – 4–5 paragraphs weaving together the Industry Challenge, Industry Shift, Our Response, Our Solution, The Difference, and Customer Proven into a persuasive, high-impact story. Use a title: The Narrative Inside the narrative, include short bold mini-headlines at the start of key sections to break up the story (e.g., “The Challenge,” “The Shift,” “The Solution”) just like in the original. End with a strong closing paragraph summarizing the opportunity and inviting engagement. 
 
Business Concept {business_concept} as a Input. 

Formatting rules:
- Always use **bold headings** exactly as written above.
- Use professional marketing tone, concise, persuasive, and confident.
- Use bullet points wherever applicable.
- Avoid markdown symbols (#, *, -) unless explicitly required for bullets.
- Final output must look like a ready-to-use B2B messaging doc.
"""
completion = client.chat.completions.create(
    model="openai/gpt-oss-120b:fireworks-ai",
    messages=[
        {
            "role": "user",
            "content": str(system_prompt)
        }
    ],
)

 

 
all_content = completion.choices[0].message.content

clean_text = re.sub(r"\*\*(.*?)\*\*", r"\1", all_content)  # remove markdown bold
clean_text = clean_text.replace("###", "").replace("##", "")
clean_text = clean_text.strip()

# === WORD DOC STRUCTURING ===
doc = Document()
lines = clean_text.split("\n")

for line in lines:
    stripped = line.strip()
    if not stripped:
        continue

    # Headings (look like section titles or ALL CAPS words)
    if re.match(r'^[A-Z][A-Za-z ]+:$', stripped) or stripped.isupper():
        doc.add_heading(stripped.rstrip(":"), level=1)

    # Subheadings (like "The Solution", "Our Promise" etc.)
    elif re.match(r'^[A-Z][A-Za-z ]+$', stripped) and len(stripped.split()) <= 5:
        doc.add_heading(stripped, level=2)

    # Bullets (starts with - or •)
    elif stripped.startswith(("- ", "• ")):
        bullet_text = stripped.lstrip("-• ").strip()
        doc.add_paragraph(bullet_text, style="List Bullet")

    # Numbered list
    elif re.match(r'^\d+\.', stripped):
        doc.add_paragraph(stripped, style="List Number")

    # Normal paragraph
    else:
        doc.add_paragraph(stripped, style="Normal")

# Save DOCX
output_path = "Professional_Brand_Messaging.docx"
doc.save(output_path)
print(f"✅ Brand messaging document saved as {output_path}")